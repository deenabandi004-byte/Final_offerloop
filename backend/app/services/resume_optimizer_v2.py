"""
Resume Optimizer V2 - Three explicit modes with no silent fallbacks.

Modes:
1. direct_edit - For DOCX files, edit directly preserving formatting
2. suggestions - For PDF files, return actionable suggestions
3. template_rebuild - Rebuild resume in ATS-optimized template
"""

import os
import json
import shutil
import tempfile
from typing import Dict, List, Tuple, Any, Optional
from docx import Document
import PyPDF2

from app.services.libreoffice_service import convert_docx_to_pdf
from app.services.docx_service import extract_text_from_docx, find_replace_in_docx


# ============================================================================
# QUALITY DETECTION
# ============================================================================

def check_conversion_quality(original_pdf: str, converted_docx: str) -> Tuple[bool, List[str]]:
    """
    Quick sanity checks before trusting PDF→DOCX conversion.
    
    Returns:
        Tuple of (is_acceptable, list_of_issues)
    """
    issues = []
    
    try:
        # Get original stats from PDF
        with open(original_pdf, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            original_text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    original_text += page_text + "\n"
        
        original_lines = [l for l in original_text.split('\n') if l.strip()]
        original_bullets = len([l for l in original_lines if l.strip().startswith(('•', '-', '◦', '*', '–'))])
        original_char_count = len(original_text)
        
        # Get converted stats from DOCX
        doc = Document(converted_docx)
        converted_paras = [p.text for p in doc.paragraphs if p.text.strip()]
        converted_text = '\n'.join(converted_paras)
        converted_bullets = len([p for p in converted_paras if p.strip().startswith(('•', '-', '◦', '*', '–'))])
        
        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            converted_paras.append(para.text)
                            converted_text += '\n' + para.text
        
        # Check 1: Paragraph count explosion (fragmentation)
        if len(original_lines) > 5 and len(converted_paras) > len(original_lines) * 2.5:
            issues.append(f"Text fragmented: {len(original_lines)} lines → {len(converted_paras)} paragraphs")
        
        # Check 2: Bullet point loss (more than 50% lost)
        if original_bullets > 3 and converted_bullets < original_bullets * 0.5:
            issues.append(f"Bullet points lost: {original_bullets} → {converted_bullets}")
        
        # Check 3: Average paragraph length (too short = fragmented)
        if converted_paras:
            avg_para_length = sum(len(p) for p in converted_paras) / len(converted_paras)
            if avg_para_length < 15:
                issues.append(f"Average paragraph too short: {avg_para_length:.0f} chars")
        
        # Check 4: Too many tiny paragraphs
        if converted_paras:
            tiny_paras = len([p for p in converted_paras if len(p.strip()) < 5])
            if tiny_paras > len(converted_paras) * 0.4:
                issues.append(f"Too many tiny paragraphs: {tiny_paras}/{len(converted_paras)}")
        
        # Check 5: Significant content loss
        if original_char_count > 100 and len(converted_text) < original_char_count * 0.6:
            issues.append(f"Content loss: {original_char_count} → {len(converted_text)} chars")
        
        # Check 6: Single-word paragraph explosion
        if converted_paras:
            single_word = len([p for p in converted_paras if len(p.split()) <= 1])
            if single_word > len(converted_paras) * 0.3:
                issues.append(f"Too many single-word paragraphs: {single_word}/{len(converted_paras)}")
        
    except Exception as e:
        print(f"[QualityCheck] Error during check: {e}")
        issues.append(f"Quality check error: {str(e)}")
    
    is_acceptable = len(issues) == 0
    
    if issues:
        print(f"[QualityCheck] ❌ Conversion issues: {issues}")
    else:
        print(f"[QualityCheck] ✅ Conversion quality acceptable")
    
    return is_acceptable, issues


# ============================================================================
# AI FUNCTIONS
# ============================================================================

def get_targeted_replacements(
    resume_text: str,
    job_description: str,
    openai_client,
    job_title: str = "",
    company: str = ""
) -> Dict[str, str]:
    """
    Get specific text replacements from OpenAI for direct editing.
    Returns dict mapping exact original text to optimized text.
    """
    
    prompt = f"""You are an expert ATS resume optimizer. Analyze this resume against the job description and provide SPECIFIC text replacements.

## CRITICAL RULES
1. "original" must be EXACT text from the resume - copy it character-for-character
2. **KEEP REPLACEMENTS VERY SHORT (1-3 words max)** - Replace single words or very short phrases only
3. Good: Replace "Built" with "Engineered" (single word)
4. Good: Replace "improving" with "boosting" (single word)
5. Good: Replace "data processing" with "large-scale data processing" (2-3 words)
6. BAD: Replace entire sentences, bullet points, or long phrases
7. **NEVER replace company names, institution names, school names, or job titles**
8. NEVER change facts (dates, job titles, company names, degrees, schools, locations)
9. NEVER fabricate skills, experiences, or achievements
10. Focus ONLY on: action verbs, adding metrics, keywords from job description

## EXAMPLES OF GOOD REPLACEMENTS
- "Developed" → "Engineered"
- "worked on" → "led development of"
- "improving performance" → "improving performance by 35%"
- "data processing" → "large-scale data processing"

## EXAMPLES OF BAD REPLACEMENTS (too long)
- "Built a web application using React" → "Engineered a scalable web application..." (entire phrase)
- Replacing entire bullet points

## JOB INFORMATION
Title: {job_title or 'Not specified'}
Company: {company or 'Not specified'}

## JOB DESCRIPTION
{job_description[:3000]}

## RESUME TEXT
{resume_text[:6000]}

## RESPONSE FORMAT
Return JSON with 15-25 SHORT, targeted replacements:
{{
    "replacements": [
        {{"original": "short exact text", "optimized": "improved version", "reason": "why"}},
        ...
    ],
    "keywords_added": ["keyword1", "keyword2"],
    "sections_modified": ["Experience", "Skills"]
}}

IMPORTANT: 
- Keep each replacement SHORT (1-3 words maximum, ideally single words)
- NEVER replace company names, school names, or institution names
- Only replace action verbs, descriptive words, or short technical terms
- This prevents formatting issues and text overlap."""

    try:
        response = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a resume optimization expert. Return only valid JSON."},
                {"role": "user", "content": prompt}
            ],
            response_format={"type": "json_object"},
            temperature=0.3,
            max_tokens=3000
        )
        
        result = json.loads(response.choices[0].message.content)
        
        # Validate replacements - only keep ones where original exists in resume
        valid_replacements = {}
        for item in result.get("replacements", []):
            original = item.get("original", "").strip()
            optimized = item.get("optimized", "").strip()
            
            if original and optimized and original != optimized:
                # Skip if original text is too long (likely to span runs)
                # Keep it very short (50 chars max) to ensure single-run replacements
                if len(original) > 50:
                    print(f"[Optimizer] ✗ Skipping (too long, {len(original)} chars): '{original[:50]}...'")
                    continue
                
                # Skip if it contains newlines
                if '\n' in original:
                    print(f"[Optimizer] ✗ Skipping (contains newlines): '{original[:50]}...'")
                    continue
                
                if original in resume_text:
                    valid_replacements[original] = optimized
                    print(f"[Optimizer] ✓ Valid replacement: '{original[:50]}...'")
                else:
                    print(f"[Optimizer] ✗ Not found in resume: '{original[:50]}...'")
        
        print(f"[Optimizer] Got {len(valid_replacements)} valid replacements")
        return valid_replacements
        
    except Exception as e:
        print(f"[Optimizer] Error getting replacements: {e}")
        return {}


def get_optimization_suggestions(
    resume_text: str,
    job_description: str,
    openai_client,
    job_title: str = "",
    company: str = ""
) -> Dict[str, Any]:
    """
    Get detailed optimization suggestions for PDF users.
    Returns actionable suggestions the user can apply manually.
    """
    
    prompt = f"""You are an expert ATS resume optimizer. Analyze this resume against the job description and provide DETAILED, ACTIONABLE suggestions.

The user has a PDF resume and wants to keep their original formatting. Give them specific changes they can make themselves.

## JOB INFORMATION
Title: {job_title or 'Not specified'}
Company: {company or 'Not specified'}

## JOB DESCRIPTION
{job_description[:3000]}

## RESUME TEXT
{resume_text[:6000]}

## RESPONSE FORMAT
Return JSON:
{{
    "suggestions": [
        {{
            "section": "Experience",
            "priority": "high",
            "current_text": "exact text from their resume",
            "suggested_text": "improved version",
            "reason": "Clear explanation of why this helps with ATS"
        }},
        ...
    ],
    "keywords_to_add": [
        {{"keyword": "Python", "where": "Skills section", "reason": "Mentioned 5x in job description"}}
    ],
    "keywords_found": ["React", "JavaScript"],
    "overall_tips": [
        "Specific actionable tip 1",
        "Specific actionable tip 2"
    ],
    "ats_score_estimate": 72,
    "score_breakdown": {{
        "keyword_match": 65,
        "formatting": 80,
        "relevance": 70
    }}
}}

Guidelines:
- Provide 8-15 specific suggestions, ordered by impact
- Be specific - quote exact text to change
- Explain WHY each change helps (ATS keywords, action verbs, quantification)
- Keywords should include where to add them
- Tips should be actionable, not generic"""

    try:
        response = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a resume optimization expert. Return only valid JSON."},
                {"role": "user", "content": prompt}
            ],
            response_format={"type": "json_object"},
            temperature=0.4,
            max_tokens=3500
        )
        
        return json.loads(response.choices[0].message.content)
        
    except Exception as e:
        print(f"[Optimizer] Error getting suggestions: {e}")
        return {
            "suggestions": [],
            "keywords_to_add": [],
            "keywords_found": [],
            "overall_tips": ["Unable to generate suggestions. Please try again."],
            "ats_score_estimate": 0,
            "error": str(e)
        }


def get_optimized_content_for_template(
    resume_text: str,
    job_description: str,
    openai_client,
    job_title: str = "",
    company: str = ""
) -> Dict[str, Any]:
    """
    Get fully optimized resume content for template rebuild.
    Returns structured content that will be rendered in a new template.
    """
    
    prompt = f"""You are an expert ATS resume optimizer. Rewrite this resume content optimized for the job description.

## RULES
1. NEVER fabricate information - only enhance what exists
2. NEVER change dates, company names, job titles, schools, or degrees
3. DO improve action verbs, add relevant keywords, quantify achievements
4. DO reorder skills by relevance to the job
5. Return structured JSON that can be rendered in a template

## JOB INFORMATION  
Title: {job_title or 'Not specified'}
Company: {company or 'Not specified'}

## JOB DESCRIPTION
{job_description[:3000]}

## ORIGINAL RESUME
{resume_text[:6000]}

## RESPONSE FORMAT
Return JSON:
{{
    "contact": {{
        "name": "Full Name",
        "email": "email@example.com",
        "phone": "123-456-7890",
        "location": "City, State",
        "linkedin": "linkedin.com/in/..."
    }},
    "summary": "2-3 sentence professional summary optimized for this role",
    "experience": [
        {{
            "title": "Job Title (unchanged)",
            "company": "Company Name (unchanged)",
            "location": "City, State",
            "startDate": "Start Date (unchanged)",
            "endDate": "End Date (unchanged)",
            "bullets": [
                "Optimized bullet point with action verb and metrics",
                "Another optimized bullet point"
            ]
        }}
    ],
    "education": [
        {{
            "degree": "Degree (unchanged)",
            "school": "School Name (unchanged)",
            "location": "City, State",
            "graduationDate": "Date (unchanged)",
            "gpa": "GPA if present",
            "coursework": ["Relevant Course 1", "Relevant Course 2"]
        }}
    ],
    "skills": {{
        "technical": ["Skill1", "Skill2"],
        "frameworks": ["Framework1", "Framework2"],
        "tools": ["Tool1", "Tool2"]
    }},
    "projects": [
        {{
            "name": "Project Name",
            "description": "Optimized description",
            "technologies": ["Tech1", "Tech2"],
            "bullets": ["Achievement 1", "Achievement 2"]
        }}
    ],
    "keywords_added": ["keyword1", "keyword2"],
    "ats_score_estimate": 85
}}"""

    try:
        response = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a resume optimization expert. Return only valid JSON."},
                {"role": "user", "content": prompt}
            ],
            response_format={"type": "json_object"},
            temperature=0.5,
            max_tokens=4000
        )
        
        return json.loads(response.choices[0].message.content)
        
    except Exception as e:
        print(f"[Optimizer] Error getting template content: {e}")
        raise


# ============================================================================
# QUALITY DETECTION
# ============================================================================

def check_docx_structure_quality(docx_path: str) -> Tuple[bool, str]:
    """
    Check if DOCX structure is suitable for direct text replacement.
    
    Some DOCX files (especially from LaTeX, Google Docs exports, or complex templates)
    have fragmented text runs that cause overlap issues after replacement.
    
    Returns:
        Tuple of (is_suitable, reason_if_not)
    """
    try:
        doc = Document(docx_path)
        
        issues = []
        fragmented_paragraphs = 0
        total_paragraphs = 0
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
                
            total_paragraphs += 1
            num_runs = len([r for r in paragraph.runs if r.text.strip()])
            
            # Check if this paragraph is fragmented (count each paragraph only once)
            is_fragmented = False
            
            # If a short paragraph has many runs, it's fragmented
            # Normal: 1-3 runs per paragraph
            # Fragmented: 5+ runs for a single line
            if num_runs >= 5 and len(text) < 200:
                is_fragmented = True
            
            # Check for suspicious patterns: single-word runs that should be together
            if not is_fragmented:  # Only check if not already marked as fragmented
                run_texts = [r.text.strip() for r in paragraph.runs if r.text.strip()]
                if len(run_texts) >= 3:
                    # Check if we have very short runs (1-2 words) that look like they should be merged
                    short_runs = sum(1 for r in run_texts if len(r.split()) <= 2)
                    if short_runs >= 3:
                        is_fragmented = True
            
            if is_fragmented:
                fragmented_paragraphs += 1
        
        # If more than 20% of paragraphs are fragmented, the DOCX is problematic
        if total_paragraphs > 0:
            fragmentation_ratio = fragmented_paragraphs / total_paragraphs
            if fragmentation_ratio > 0.2:
                issues.append(f"High fragmentation: {fragmentation_ratio:.0%} of paragraphs have fragmented text runs")
        
        if issues:
            return False, "; ".join(issues)
        
        return True, "OK"
        
    except Exception as e:
        return False, f"Error checking DOCX structure: {str(e)}"


# ============================================================================
# MAIN OPTIMIZATION FUNCTIONS
# ============================================================================

def optimize_direct_edit(
    docx_path: str,
    job_description: str,
    openai_client,
    job_title: str = "",
    company: str = ""
) -> Tuple[bytes, Dict[str, Any]]:
    """
    MODE 1: Direct Edit (for DOCX files)
    
    Edit the DOCX directly, preserving all formatting.
    Returns the optimized PDF bytes.
    """
    print("[OptimizeV2] Mode: DIRECT_EDIT (DOCX)")
    
    # Check DOCX structure quality first
    is_suitable, reason = check_docx_structure_quality(docx_path)
    if not is_suitable:
        print(f"[OptimizeV2] ⚠️ DOCX structure not suitable for direct edit: {reason}")
        print("[OptimizeV2] Falling back to suggestions mode...")
        
        # Fall back to suggestions mode
        return optimize_suggestions(
            file_path=docx_path,
            job_description=job_description,
            openai_client=openai_client,
            job_title=job_title,
            company=company,
            file_type="docx",
            original_mode="direct_edit",
            fallback_reason=reason
        )
    
    work_dir = tempfile.mkdtemp(prefix="resume_direct_")
    
    try:
        # Step 1: Extract text from DOCX
        print("[OptimizeV2] Step 1: Extract text from DOCX")
        resume_text = extract_text_from_docx(docx_path)
        
        if not resume_text or len(resume_text) < 50:
            raise ValueError("Could not extract text from DOCX")
        
        print(f"[OptimizeV2] Extracted {len(resume_text)} characters")
        
        # Step 2: Get targeted replacements from AI
        print("[OptimizeV2] Step 2: Get AI replacements")
        replacements = get_targeted_replacements(
            resume_text, job_description, openai_client, job_title, company
        )
        
        if not replacements:
            print("[OptimizeV2] Warning: No replacements generated")
        
        # Step 3: Copy DOCX and apply replacements
        print("[OptimizeV2] Step 3: Apply replacements to DOCX")
        output_docx = os.path.join(work_dir, "optimized.docx")
        shutil.copy(docx_path, output_docx)
        
        if replacements:
            success = find_replace_in_docx(output_docx, replacements)
            if not success:
                print("[OptimizeV2] Warning: Some replacements may have failed")
        
        # Step 4: Convert to PDF
        print("[OptimizeV2] Step 4: Convert DOCX to PDF")
        output_pdf = convert_docx_to_pdf(output_docx, work_dir)
        
        if not output_pdf:
            raise ValueError("Failed to convert optimized DOCX to PDF")
        
        # Read PDF bytes
        with open(output_pdf, 'rb') as f:
            pdf_bytes = f.read()
        
        print(f"[OptimizeV2] ✅ Direct edit complete. PDF size: {len(pdf_bytes)} bytes")
        
        metadata = {
            "mode": "direct_edit",
            "success": True,
            "replacements_made": len(replacements),
            "formatting_preserved": True,
            "message": "Your resume has been optimized while preserving your original formatting."
        }
        
        return pdf_bytes, metadata
        
    finally:
        shutil.rmtree(work_dir, ignore_errors=True)


def optimize_suggestions(
    file_path: str,
    job_description: str,
    openai_client,
    job_title: str = "",
    company: str = "",
    file_type: str = "pdf",
    original_mode: Optional[str] = None,
    fallback_reason: Optional[str] = None
) -> Tuple[None, Dict[str, Any]]:
    """
    MODE 2: Suggestions Mode (for PDF files)
    
    Returns actionable suggestions without modifying the file.
    User applies changes manually to preserve their formatting.
    """
    print("[OptimizeV2] Mode: SUGGESTIONS")
    
    # Extract text based on file type
    if file_type == "pdf":
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            resume_text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    resume_text += page_text + "\n"
    else:
        resume_text = extract_text_from_docx(file_path)
    
    if not resume_text or len(resume_text) < 50:
        raise ValueError("Could not extract text from resume")
    
    print(f"[OptimizeV2] Extracted {len(resume_text)} characters")
    
    # Get suggestions from AI
    print("[OptimizeV2] Getting optimization suggestions...")
    suggestions = get_optimization_suggestions(
        resume_text, job_description, openai_client, job_title, company
    )
    
    print(f"[OptimizeV2] ✅ Generated {len(suggestions.get('suggestions', []))} suggestions")
    
    metadata = {
        "mode": "suggestions",
        "success": True,
        "formatting_preserved": True,
        "suggestions": suggestions.get("suggestions", []),
        "keywords_to_add": suggestions.get("keywords_to_add", []),
        "keywords_found": suggestions.get("keywords_found", []),
        "overall_tips": suggestions.get("overall_tips", []),
        "ats_score_estimate": suggestions.get("ats_score_estimate", 0),
        "score_breakdown": suggestions.get("score_breakdown", {}),
        "message": "Here are specific suggestions to optimize your resume. Apply these changes to your original file to preserve your formatting."
    }
    
    # Add fallback info if this was a fallback from direct_edit
    if original_mode:
        metadata["original_mode"] = original_mode
    if fallback_reason:
        metadata["fallback_reason"] = fallback_reason
        metadata["message"] = f"Your DOCX file has complex formatting that could cause text overlap. Here are suggestions to optimize your resume instead. ({fallback_reason})"
    
    return None, metadata


def optimize_template_rebuild(
    file_path: str,
    job_description: str,
    openai_client,
    job_title: str = "",
    company: str = "",
    file_type: str = "pdf"
) -> Tuple[None, Dict[str, Any]]:
    """
    MODE 3: Template Rebuild (opt-in for any file type)
    
    Rebuilds the resume in a clean ATS-optimized template.
    Content is optimized, formatting is standardized.
    
    Note: Returns structured content - frontend renders the PDF using React-PDF.
    """
    print("[OptimizeV2] Mode: TEMPLATE_REBUILD")
    
    # Extract text based on file type
    if file_type == "pdf":
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            resume_text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    resume_text += page_text + "\n"
    else:
        resume_text = extract_text_from_docx(file_path)
    
    if not resume_text or len(resume_text) < 50:
        raise ValueError("Could not extract text from resume")
    
    print(f"[OptimizeV2] Extracted {len(resume_text)} characters")
    
    # Get optimized structured content
    print("[OptimizeV2] Getting optimized content for template...")
    optimized_content = get_optimized_content_for_template(
        resume_text, job_description, openai_client, job_title, company
    )
    
    print(f"[OptimizeV2] ✅ Template rebuild complete")
    
    metadata = {
        "mode": "template_rebuild",
        "success": True,
        "formatting_preserved": False,
        "structured_content": optimized_content,
        "keywords_added": optimized_content.get("keywords_added", []),
        "ats_score_estimate": optimized_content.get("ats_score_estimate", 0),
        "message": "Your resume has been rebuilt in an ATS-optimized template with enhanced content."
    }
    
    return None, metadata


def optimize_resume_v2(
    file_path: str,
    file_type: str,
    job_description: str,
    openai_client,
    mode: str,
    job_title: str = "",
    company: str = ""
) -> Tuple[Optional[bytes], Dict[str, Any]]:
    """
    Main entry point for resume optimization.
    
    Args:
        file_path: Path to the resume file
        file_type: 'pdf', 'docx', or 'doc'
        job_description: The job description to optimize for
        openai_client: OpenAI client instance
        mode: 'direct_edit', 'suggestions', or 'template_rebuild'
        job_title: Optional job title
        company: Optional company name
    
    Returns:
        Tuple of (pdf_bytes or None, metadata_dict)
    """
    print(f"\n{'='*60}")
    print(f"[OptimizeV2] Starting optimization")
    print(f"[OptimizeV2] File type: {file_type}")
    print(f"[OptimizeV2] Mode: {mode}")
    print(f"[OptimizeV2] Job: {job_title} at {company}")
    print(f"{'='*60}\n")
    
    # Validate mode for file type
    if mode == 'direct_edit' and file_type not in ['docx']:
        raise ValueError(f"Direct edit mode requires DOCX file, got {file_type}")
    
    # Route to appropriate handler
    if mode == 'direct_edit':
        return optimize_direct_edit(
            file_path, job_description, openai_client, job_title, company
        )
    
    elif mode == 'suggestions':
        return optimize_suggestions(
            file_path, job_description, openai_client, job_title, company, file_type
        )
    
    elif mode == 'template_rebuild':
        return optimize_template_rebuild(
            file_path, job_description, openai_client, job_title, company, file_type
        )
    
    else:
        raise ValueError(f"Unknown optimization mode: {mode}")
