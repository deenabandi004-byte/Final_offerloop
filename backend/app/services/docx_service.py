"""
DOCX manipulation service for text replacement while preserving formatting
"""
from docx import Document
from typing import Dict, Optional


def extract_text_from_docx(docx_path: str) -> Optional[str]:
    """
    Extract plain text from DOCX file.
    
    Args:
        docx_path: Path to DOCX file
    
    Returns:
        Plain text content, or None if extraction failed
    """
    try:
        doc = Document(docx_path)
        text_parts = []
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(paragraph.text)
        
        return '\n'.join(text_parts)
        
    except Exception as e:
        print(f"[DOCX] Error extracting text from DOCX: {e}")
        import traceback
        traceback.print_exc()
        return None


def find_replace_in_docx(
    docx_path: str,
    replacements: Dict[str, str],
    output_path: Optional[str] = None
) -> bool:
    """
    Find and replace text in DOCX while preserving formatting.
    
    CRITICAL: Only replaces text found entirely within a SINGLE run.
    This prevents text overlap issues that occur when text spans multiple runs.
    
    Args:
        docx_path: Path to input DOCX file
        replacements: Dictionary mapping old text to new text
        output_path: Path for output file (defaults to overwriting input)
    
    Returns:
        True if successful, False otherwise
    """
    try:
        if not replacements:
            print("[DOCX] No replacements to make")
            return True
        
        doc = Document(docx_path)
        replacements_made = 0
        
        # Process paragraphs - ONLY replace text in single runs
        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                # Only replace if text exists entirely within a SINGLE run
                for run in paragraph.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
                        replacements_made += 1
                        print(f"[DOCX] ✓ Replaced in single run: '{old_text[:40]}...'")
                        break  # Only replace once per paragraph
        
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in replacements.items():
                            # Only replace if text exists entirely within a SINGLE run
                            for run in paragraph.runs:
                                if old_text in run.text:
                                    run.text = run.text.replace(old_text, new_text)
                                    replacements_made += 1
                                    break  # Only replace once per paragraph
        
        # Save
        output = output_path or docx_path
        doc.save(output)
        print(f"[DOCX] ✅ Made {replacements_made} replacements (single-run only), saved to {output}")
        return True
        
    except Exception as e:
        print(f"[DOCX] Error replacing text in DOCX: {e}")
        import traceback
        traceback.print_exc()
        return False
